# exporters.py
from docx import Document
from docx.shared import Pt
from pathlib import Path
from typing import List, Dict

def md_chapters_to_docx(chapters: List[Dict], references: Dict[str,str], out_path: Path) -> Path:
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    for ch in chapters:
        doc.add_heading(ch['title'], level=1)
        for sec in ch.get('sections', []):
            doc.add_heading(sec['heading'], level=2)
            for para in sec.get('paragraphs', []):
                doc.add_paragraph(para)

    doc.add_heading("References", level=1)
    for key, ref in references.items():
        doc.add_paragraph(f"{key}: {ref}")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return out_path
