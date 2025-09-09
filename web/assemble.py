from __future__ import annotations
from pathlib import Path
from typing import List, Dict
import json

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- tiny helpers ------------------------------------------------------------

def _read_text(path: Path) -> str:
    return Path(path).read_text(encoding="utf-8")

def _read_sources(path: Path) -> List[Dict]:
    try:
        data = json.loads(Path(path).read_text(encoding="utf-8"))
        # normalize: ensure keys exist and sort by key like S1, S2, ...
        for i, s in enumerate(data, 1):
            s.setdefault("key", f"S{i}")
            s.setdefault("title", "Untitled")
            s.setdefault("year", None)
            s.setdefault("doi", None)
            s.setdefault("url", None)
        data.sort(key=lambda s: (s.get("key") or "S9999"))
        return data
    except Exception:
        return []

def _fmt_ref(s: Dict) -> str:
    # Simple APA-ish line: [S1] Title. (Year) DOI/URL
    parts = [f"[{s.get('key','S?')}] {s.get('title','Untitled')}."]
    if s.get("year"):
        parts.append(f"({s['year']})")
    if s.get("doi"):
        parts.append(f"https://doi.org/{s['doi']}")
    elif s.get("url"):
        parts.append(s["url"])
    return " ".join(parts)

# --- very light Markdown-to-DOCX --------------------------------------------

def _add_paragraph(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p.style = doc.styles["Normal"]
    return p

def _add_heading(doc: Document, text: str, level: int):
    # python-docx: 0=Title, 1..9 heading levels. We'll use 1..3.
    h = doc.add_heading(text, level=level)
    return h

def _add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(text, style="List Bullet")
    return p

def _init_styles(doc: Document):
    # make Normal 11pt for readability
    normal = doc.styles["Normal"]
    if normal.font.size is None or normal.font.size.pt != 11:
        normal.font.size = Pt(11)

def _parse_markdown_into_doc(doc: Document, md_text: str):
    """
    Minimal parser:
    - '# ' → Heading 1, '## ' → Heading 2, '### ' → Heading 3
    - '- ' → bullet list items
    - blank line separates paragraphs
    - everything else is a paragraph
    """
    lines = md_text.splitlines()
    buf: List[str] = []
    list_mode = False

    def flush_buf():
        nonlocal buf
        if not buf:
            return
        _add_paragraph(doc, " ".join(buf).strip())
        buf = []

    for raw in lines:
        line = raw.rstrip()

        # headings
        if line.startswith("### "):
            flush_buf()
            _add_heading(doc, line[4:].strip(), level=3)
            list_mode = False
            continue
        if line.startswith("## "):
            flush_buf()
            _add_heading(doc, line[3:].strip(), level=2)
            list_mode = False
            continue
        if line.startswith("# "):
            flush_buf()
            _add_heading(doc, line[2:].strip(), level=1)
            list_mode = False
            continue

        # bullets
        if line.lstrip().startswith("- "):
            flush_buf()
            _add_bullet(doc, line.lstrip()[2:].strip())
            list_mode = True
            continue

        # blank → new paragraph
        if not line.strip():
            flush_buf()
            list_mode = False
            continue

        # normal text line
        if list_mode:
            # treat as continuation of a bullet block: add paragraph
            _add_paragraph(doc, line.strip())
        else:
            buf.append(line.strip())

    flush_buf()

def assemble_to_docx(final_md_path: Path, sources_json_path: Path, out_docx_path: Path) -> Path:
    """
    Build a DOCX from a markdown body (final_md_path) and a sources.json list.
    Returns the output path.
    """
    final_md_path = Path(final_md_path)
    sources_json_path = Path(sources_json_path)
    out_docx_path = Path(out_docx_path)
    out_docx_path.parent.mkdir(parents=True, exist_ok=True)

    body_md = _read_text(final_md_path)
    sources = _read_sources(sources_json_path)

    doc = Document()
    _init_styles(doc)

    # Optional title (use filename stem)
    title = final_md_path.stem.replace("_", " ").title()
    title_para = doc.add_paragraph(title)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.runs[0].bold = True

    doc.add_paragraph("")  # spacer

    # Body
    _parse_markdown_into_doc(doc, body_md)

    # References
    if sources:
        doc.add_page_break()
        _add_heading(doc, "References", level=1)
        for s in sources:
            _add_paragraph(doc, _fmt_ref(s))

    doc.save(out_docx_path)
    return out_docx_path
